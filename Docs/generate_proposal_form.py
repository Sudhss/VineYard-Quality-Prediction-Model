import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_centered_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold)

def add_left_text(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold)

def create_proposal_form():
    doc = Document()

    # --- Header Page 1 ---
    add_left_text(doc, "Annexure-1\n", size=10, bold=True)
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Emerging Technologies\n", size=12, bold=True)
    add_left_text(doc, "Page 1 of 2", size=10)
    
    add_centered_text(doc, "PROJECT TITLE AND SUPERVISOR PROPOSAL FORM\n", size=13, bold=True)

    # --- Form Details ---
    add_left_text(doc, "Project Group No. (Allotted by Project Coordinator): GP 42")
    add_left_text(doc, "Year of Student: 3rd Year\t\t\tSemester: VI")
    add_left_text(doc, "Propose Project Title: Vineyard Quality Assesment Model : using Machine Learning.")
    add_left_text(doc, "Project Domain: Artificial Intelligence & Machine Learning - Predictive Analysis.")
    add_left_text(doc, "Name of Student (Group Leader): Sparsh Bhalla")
    
    doc.add_paragraph("\n")
    add_left_text(doc, "Details of Group Members:", bold=True)

    # --- Table of Members ---
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.\nNo.", "Roll No. of Student", "Name of Students", "Section", "Signature of\nStudents"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    members = [
        ("1", "2301331530173", "Sparsh Bhalla", "C", ""),
        ("2", "2301331530176", "Sudhanshu Shukla", "C", ""),
        ("3", "2301331530174", "Srijan Yadav", "C", ""),
        ("4", "2301331530175", "Sriom Kant", "C", "")
    ]

    for row_idx, member in enumerate(members, start=1):
        for col_idx, text in enumerate(member):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    add_left_text(doc, "Project Type (Tick one):", bold=True)
    add_left_text(doc, "[X] Mini Project\t\t[ ] Minor Project\t\t[ ] Major Project")
    doc.add_paragraph("\n")

    # --- Problem Statement ---
    add_left_text(doc, "Brief About Project (Problem Statement) (Max 250 words):", bold=True)
    brief = ("Traditional wine quality assessment relies on slow, subjective human sensory evaluation. "
             "This project solves this by developing an XGBoost machine learning model that accurately "
             "predicts wine quality using 11 physicochemical properties. The model is deployed via a Streamlit "
             "web application, providing instant, data-driven agricultural insights to vineyard owners.")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(brief)
    set_font(run, 11)
    
    doc.add_paragraph("\n")

    # --- Proposed Objectives ---
    add_left_text(doc, "Proposed Objectives:", bold=True)
    add_left_text(doc, "1. Train a highly accurate XGBoost machine learning model on physicochemical tabular data.")
    add_left_text(doc, "2. Develop a responsive Streamlit web application for real-time quality score predictions.")
    add_left_text(doc, "3. Implement feature explainability to provide vineyard owners with actionable insights.")

    doc.add_page_break()

    # --- Page 2 ---
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Emerging Technologies\n", size=12, bold=True)
    add_left_text(doc, "Page 2 of 2", size=10)
    doc.add_paragraph("\n")

    add_left_text(doc, "Expected Outcomes (Tick all that apply):", bold=True)
    add_left_text(doc, "[ ] Research Paper Publication (SCI/Scopus)\t\t[ ] Patent Filing")
    add_left_text(doc, "[X] Product Development\t\t\t\t[ ] Startup Idea")
    add_left_text(doc, "[ ] SIH/Industry Collaboration\t\t\t[ ] Others")
    doc.add_paragraph("\n")

    add_left_text(doc, "Proposed Supervisor Name: Mr. Ritesh Rajput")
    add_left_text(doc, "Propose Co-Supervisor Name (if required): ________________________________________________")
    doc.add_paragraph("\n\n\n")

    add_left_text(doc, "__________________________\nName and Signature\nSupervisor(s)")
    doc.add_paragraph("\n")

    add_left_text(doc, "Remark by Project Coordinator/Co-Coordinator/PCEC", bold=True)
    add_left_text(doc, "Project Committee Review Comments:", bold=True)
    review = "The proposal is highly relevant to modern agricultural technology. Using XGBoost for tabular data is an excellent technical choice, and wrapping it in a Streamlit UI ensures the project is a complete 'Product Development' outcome rather than just a standalone script. Approved to proceed."
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run2 = p2.add_run(review)
    set_font(run2, 11)
    
    doc.add_paragraph("\n")
    add_left_text(doc, "Project Approval Status:", bold=True)
    add_left_text(doc, "[X] Approved\t\t[ ] Revision Required\t\t[ ] Reject")
    doc.add_paragraph("\n\n\n")

    add_left_text(doc, "_____________________\t\t_____________________\t\t_____________________")
    add_left_text(doc, "Supervisor(s)\t\t\tProject Coordinator/Co-Coordinator\tPCEC")
    doc.add_paragraph("\n")
    add_left_text(doc, "Date of Approval: ________________")

    # Save Document
    doc.save("Annexure1_Proposal_Form.docx")
    print(f"Generated successfully: Annexure1_Proposal_Form.docx")

if __name__ == "__main__":
    create_proposal_form()
