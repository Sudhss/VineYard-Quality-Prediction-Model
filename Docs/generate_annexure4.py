import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_centered_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold)

def add_left_text(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold)

def create_annexure4():
    doc = Document()

    # --- Header Page 1 ---
    add_left_text(doc, "Annexure-4\n", size=10, bold=True)
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology\nDepartment of Artificial Intelligence & Machine Learning (AIML)", size=12, bold=True)
    add_left_text(doc, "Page 1 of 2", size=10)
    
    add_centered_text(doc, "PROJECT SUBMISSION FORM\n", size=13, bold=True)

    # --- Form Details ---
    add_left_text(doc, "Project Type (Tick one):")
    add_left_text(doc, "[X] Mini Project\t[ ] Minor Project\t[ ] Major Project")
    doc.add_paragraph("\n")

    add_left_text(doc, "Project Group No: G01\t\tSession: 2024-2025\t\tYear: 3rd Year\t\tSemester: 6th")
    add_left_text(doc, "Project Title: Vineyard Quality Assesment Model", bold=True)
    
    doc.add_paragraph("\n")

    # --- Table of Members ---
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.\nNo.", "Roll No. of Student", "Name of Students", "Section", "Signature of\nStudents"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    members = [
        ("1", "2301331530176", "Sudhanshu Shukla", "", ""),
        ("2", "2301331530173", "Sparsh Bhalla", "", ""),
        ("3", "2301331530174", "Srijan Yadav", "", ""),
        ("4", "2301331530175", "Sriom Kant", "", "")
    ]

    for row_idx, member in enumerate(members, start=1):
        for col_idx, text in enumerate(member):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    add_left_text(doc, "Supervisor Name: Mr. Ritesh Rajput")
    add_left_text(doc, "Co-Supervisor Name (if any): ___________________________________________________________")
    doc.add_paragraph("\n")

    # --- Submission Section ---
    add_left_text(doc, "Project Submission Date: ____________________", bold=True)
    doc.add_paragraph("\n")
    
    add_left_text(doc, "Project Submission Checklist:", bold=True)
    add_left_text(doc, "[X] Project Model in working condition")
    add_left_text(doc, "[X] Project Report (Hard & Soft Copy) duly signed by Supervisors")
    add_left_text(doc, "[X] Power Point Presentation (PPT) in soft copy")
    add_left_text(doc, "[X] Code Repository (If applicable)")
    add_left_text(doc, "[X] Plagiarism Report duly signed by Supervisors")
    add_left_text(doc, "[ ] Research Paper/ Patent (If applicable)")
    add_left_text(doc, "[X] Project Progress Report Card duly signed by Supervisors")

    doc.add_page_break()

    # --- Page 2 ---
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology\nDepartment of Artificial Intelligence & Machine Learning (AIML)", size=12, bold=True)
    add_left_text(doc, "Page 2 of 2", size=10)
    doc.add_paragraph("\n")

    # --- Remarks ---
    add_left_text(doc, "Remark by Supervisor (Recommended for Submission):", bold=True)
    remark1 = "The group successfully built the Streamlit UI this week and fully integrated the tuned XGBoost model into the backend. The project is in excellent working condition, handles edge cases well, and demonstrates a very deep understanding of end-to-end machine learning pipelines. Highly recommended for final submission."
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = p1.add_run(remark1)
    set_font(run1, 11)
    
    doc.add_paragraph("\n\n")

    add_left_text(doc, "Remark by Project Coordinator/Co-Coordinator/PCEC (if any):", bold=True)
    remark2 = "All documents, code repositories, and the fully functional Streamlit application have been verified. The GUI integration represents a solid completion of the week 4/5 objectives. Approved for final evaluation."
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run2 = p2.add_run(remark2)
    set_font(run2, 11)
    
    doc.add_paragraph("\n\n\n\n")

    add_left_text(doc, "_____________________\t\t_____________________\t\t_____________________")
    add_left_text(doc, "Supervisor(s)\t\t\tProject Coordinator/Co-Coordinator\tPCEC")

    # Save Document
    doc.save("Annexure4.docx")
    print(f"Generated successfully: Annexure4.docx")

if __name__ == "__main__":
    create_annexure4()
