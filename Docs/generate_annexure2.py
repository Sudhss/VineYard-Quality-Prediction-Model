import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_centered_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold)

def add_left_text(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold)

def create_annexure2():
    doc = Document()

    # --- Header Page 1 ---
    add_left_text(doc, "Annexure-2\n", size=10, bold=True)
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Emerging Technologies\n", size=12, bold=True)
    add_left_text(doc, "Page 1 of 2", size=10)
    
    add_centered_text(doc, "PROJECT PROGRESS EVALUATION REPORT FORMAT\n(Review-3)", size=13, bold=True)

    # --- Form Details ---
    add_left_text(doc, "Project Type (Tick one):")
    add_left_text(doc, "[X] Mini Project\t[ ] Minor Project\t[ ] Major Project")
    doc.add_paragraph("\n")

    add_left_text(doc, "Project Group No: GP 42\t\tSession: 2024-2025\t\tYear: 3rd Year\t\tSemester: VI")
    add_left_text(doc, "Project Title: Vineyard Quality Assesment Model : using Machine Learning.", bold=True)
    
    doc.add_paragraph("\n")

    # --- Table of Members ---
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.\nNo.", "Roll No. of Student", "Name of Students", "Section", "Signature of\nStudents"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    members = [
        ("1", "2301331530173", "Sparsh Bhalla", "C", ""),
        ("2", "2301331530176", "Sudhanshu Shukla", "C", ""),
        ("3", "2301331530174", "Srijan Yadav", "C", ""),
        ("4", "2301331530175", "Sriom Kant", "C", "")
    ]

    for row_idx, member in enumerate(members, start=1):
        for col_idx, text in enumerate(member):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    add_left_text(doc, "Supervisor Name: Mr. Ritesh Rajput")
    add_left_text(doc, "Co-Supervisor Name (if any): ___________________________________________________________")
    doc.add_paragraph("\n")

    # --- Supervisor Section ---
    add_left_text(doc, "To be filled by Supervisors", bold=True)
    add_left_text(doc, "I have seen the progress report: [X] Yes / [ ] No")
    add_left_text(doc, "I have seen the PPT: [X] Yes / [ ] No")
    add_left_text(doc, "Recommended for presentation: [X] Yes")
    doc.add_paragraph("\n")
    
    add_left_text(doc, "Remark, if any:", bold=True)
    remark = "The group has successfully developed and deployed the XGBoost Machine Learning model. The Streamlit web interface is fully functional and accurately predicts wine quality based on chemical inputs. Excellent overall execution of the final project."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(remark)
    set_font(run, 11)

    doc.add_paragraph("\n\n")
    add_left_text(doc, "__________________________\nName and Signature\nSupervisor")
    
    doc.add_paragraph("\n")
    
    # --- Committee Section ---
    add_left_text(doc, "Remark by Project Coordinator/Co-Coordinator/PCEC", bold=True)
    add_left_text(doc, "PCEC Review Comments:", bold=True)
    pcec = "The project demonstrates a solid understanding of predictive analytics. The integration of the tuned XGBoost model with a user-friendly UI fulfills all proposed project objectives. Highly satisfactory."
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run2 = p2.add_run(pcec)
    set_font(run2, 11)

    doc.add_page_break()

    # --- Page 2 ---
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Emerging Technologies\n", size=12, bold=True)
    add_left_text(doc, "Page 2 of 2", size=10)
    doc.add_paragraph("\n")

    # --- Table of Grades ---
    table2 = doc.add_table(rows=5, cols=5)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ["Sr.\nNo.", "Roll No. of Student", "Name of Students", "Section", "Evaluation Grade\n(As per rubrics)"]
    for i, header in enumerate(headers2):
        cell = table2.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    grades = [
        ("1", "2301331530173", "Sparsh Bhalla", "C", "S"),
        ("2", "2301331530176", "Sudhanshu Shukla", "C", "S"),
        ("3", "2301331530174", "Srijan Yadav", "C", "S"),
        ("4", "2301331530175", "Sriom Kant", "C", "S")
    ]

    for row_idx, member in enumerate(grades, start=1):
        for col_idx, text in enumerate(member):
            cell = table2.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    add_left_text(doc, "Overall progress of Group:", bold=True)
    add_left_text(doc, "[X] Excellent (S)\t[ ] Good (A)\t[ ] Satisfactory (B)\t[ ] Unsatisfactory (C)")
    
    doc.add_paragraph("\n\n\n")

    add_left_text(doc, "_____________________\t\t_____________________\t\t_____________________")
    add_left_text(doc, "Supervisor(s)\t\t\tProject Coordinator/Co-Coordinator\tPCEC")
    doc.add_paragraph("\n")
    add_left_text(doc, "Date of Evaluation: ________________")

    # Save Document
    doc.save("Annexure2_Evaluation_Report.docx")
    print(f"Generated successfully: Annexure2_Evaluation_Report.docx")

if __name__ == "__main__":
    create_annexure2()
