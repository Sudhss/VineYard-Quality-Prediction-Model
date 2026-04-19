import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_centered_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold)

def add_left_text(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold)

def create_annexure1():
    doc = Document()

    # --- Header Page 1 ---
    add_left_text(doc, "Annexure-1\n", size=10, bold=True)
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology\nDepartment of Artificial Intelligence & Machine Learning (AIML)", size=12, bold=True)
    add_left_text(doc, "Page 1 of 2", size=10)
    
    add_centered_text(doc, "PROJECT TITLE AND SUPERVISOR PROPOSAL FORM\n", size=13, bold=True)

    # --- Form Details ---
    add_left_text(doc, "Project Group No. (Allotted by Project Coordinator): ____________________")
    add_left_text(doc, "Year of Student: 3rd Year\t\t\tSemester: 6th")
    add_left_text(doc, "Propose Project Title: Vineyard Quality Assesment Model", bold=True)
    add_left_text(doc, "Project Domain: Machine Learning / Artificial Intelligence")
    add_left_text(doc, "Name of Student (Group Leader): Sudhanshu Shukla")
    
    doc.add_paragraph("\n")
    add_left_text(doc, "Details of Group Members:", bold=True)

    # --- Table of Members ---
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Sr.\nNo.", "Roll No. of Student", "Name of Students", "Section", "Signature of\nStudents"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    members = [
        ("1", "2301331530176", "Sudhanshu Shukla", "", ""),
        ("2", "2301331530173", "Sparsh Bhalla", "", ""),
        ("3", "2301331530174", "Srijan Yadav", "", ""),
        ("4", "2301331530175", "Sriom Kant", "", "")
    ]

    for row_idx, member in enumerate(members, start=1):
        for col_idx, text in enumerate(member):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    # --- Project Details ---
    add_left_text(doc, "Project Type (Tick one):")
    add_left_text(doc, "[X] Mini Project\t[ ] Minor Project\t[ ] Major Project")
    doc.add_paragraph("\n")

    add_left_text(doc, "Brief About Project (Problem Statement) (Max 250 words):", bold=True)
    prob_statement = "The viticulture industry currently relies heavily on manual, sensory evaluation to determine the quality of wine yields, a process that is time-consuming and prone to human error. To automate this, our project proposes an initial Machine Learning system to predict wine quality purely based on physicochemical properties (such as pH, alcohol content, and acidity). For this phase, we have ingested the UCI Wine Quality Dataset and established a baseline prediction capability using a Random Forest regression model. Random Forest was chosen as the initial baseline algorithm due to its simplicity and resistance to overfitting on tabular data. By analyzing multiple decision trees, we aim to uncover basic correlations between chemical compounds and overall wine quality, laying the groundwork for a more advanced prediction pipeline."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(prob_statement)
    set_font(run, 11)
    
    doc.add_paragraph("\n")

    add_left_text(doc, "Proposed Objectives:", bold=True)
    objs = [
        "1. Ingest and perform exploratory data analysis on the UCI Wine Quality Dataset.",
        "2. Train and evaluate a baseline Random Forest regression model.",
        "3. Identify key chemical correlations impacting vineyard yield quality."
    ]
    for obj in objs:
        add_left_text(doc, obj)

    doc.add_page_break()

    # --- Page 2 ---
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology\nDepartment of Artificial Intelligence & Machine Learning (AIML)", size=12, bold=True)
    add_left_text(doc, "Page 2 of 2", size=10)
    doc.add_paragraph("\n")

    add_left_text(doc, "Expected Outcomes (Tick all that apply):", bold=True)
    add_left_text(doc, "[ ] Research Paper Publication (SCI/Scopus)\t[ ] Patent Filing")
    add_left_text(doc, "[X] Product Development\t\t\t[ ] Startup Idea")
    add_left_text(doc, "[ ] SIH/Industry Collaboration\t\t\t[ ] Others")
    
    doc.add_paragraph("\n")
    add_left_text(doc, "Proposed Supervisor Name: Mr. Ritesh Rajput", bold=True)
    add_left_text(doc, "Propose Co-Supervisor Name (if required): ________________________________________________")
    
    doc.add_paragraph("\n\n\n")
    
    # Signatures
    add_left_text(doc, "__________________________\t\t\t\t\t__________________________")
    add_left_text(doc, "Name and Signature\t\t\t\t\t\tSupervisor(s)")
    
    doc.add_paragraph("\n")
    
    # Approvals
    add_left_text(doc, "Remark by Project Coordinator/Co-Coordinator/PCEC", bold=True)
    doc.add_paragraph("\n")
    
    add_left_text(doc, "Project Committee Review Comments:", bold=True)
    add_left_text(doc, "Baseline Random Forest model provides a decent starting point, but it lacks the precision and ")
    add_left_text(doc, "feature explainability required for complex chemical tabular data. For your next iteration, ")
    add_left_text(doc, "you must switch the core algorithm to Extreme Gradient Boosting (XGBoost).")
    doc.add_paragraph("\n")
    
    add_left_text(doc, "Project Approval Status:", bold=True)
    add_left_text(doc, "[X] Approved\t\t[ ] Revision Required\t\t[ ] Reject")
    doc.add_paragraph("\n\n\n")

    add_left_text(doc, "__________________________\t\t\t\t\t__________________________")
    add_left_text(doc, "Supervisor(s)\t\t\t\t\t\tProject Coordinator/Co-Coordinator PCEC")
    doc.add_paragraph("\n")
    add_left_text(doc, "Date of Approval: ________________")

    # Save Document
    # Saving directly to the current directory (which will be Docs)
    doc.save("Annexure1.docx")
    print(f"Generated successfully: Annexure1.docx")

if __name__ == "__main__":
    create_annexure1()
