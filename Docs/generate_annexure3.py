import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_centered_text(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold)

def add_left_text(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size, bold)

def create_annexure3():
    doc = Document()

    # --- Header Page 1 ---
    add_left_text(doc, "Annexure-3\n", size=10, bold=True)
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Gr. Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology\nDepartment of Artificial Intelligence & Machine Learning (AIML)", size=12, bold=True)
    
    add_centered_text(doc, "\nPROJECT PROGRESS REPORT CARD", size=13, bold=True)
    add_centered_text(doc, "Session: 2024 - 2025\n", size=11, bold=True)

    # --- Form Details ---
    add_left_text(doc, "Project Type (Tick one):")
    add_left_text(doc, "[X] Mini Project\t[ ] Minor Project\t[ ] Major Project")
    doc.add_paragraph("\n")

    add_left_text(doc, "Student Details:", bold=True)
    add_left_text(doc, "Name: Sudhanshu Shukla (Group Leader)")
    add_left_text(doc, "Roll Number: 2301331530176")
    add_left_text(doc, "Year: 3rd Year")
    add_left_text(doc, "Semester: 6th")
    add_left_text(doc, "Section: ____________")
    add_left_text(doc, "Project Group No.: G01")
    add_left_text(doc, "Project Title: Vineyard Quality Assesment Model", bold=True)
    add_left_text(doc, "Supervisor Name: Mr. Ritesh Rajput")
    add_left_text(doc, "Co-Supervisor Name (if any): _________________________________")
    add_left_text(doc, "Student Contact No.: _________________________________________")
    
    doc.add_paragraph("\n")
    
    add_left_text(doc, "Student Declaration:", bold=True)
    add_left_text(doc, "I, Sudhanshu Shukla, hereby declare that the above progress details are accurate and updated as per my project work.")
    add_left_text(doc, "Student's Signature: _____________________________")
    doc.add_paragraph("\n")

    add_left_text(doc, "Final Remarks by Project Coordinator/Co-coordinator:", bold=True)
    p = doc.add_paragraph()
    run = p.add_run("Consistent and excellent progress shown across the first three weeks. The logical progression from a baseline Random Forest model to a fully fine-tuned XGBoost architecture using GridSearchCV is well documented and scientifically sound. Model weights are now finalized. Approved to proceed to the Streamlit UI development phase.")
    set_font(run, 11)
    
    doc.add_paragraph("\n\n\n")
    add_left_text(doc, "__________________________\t\t\t\t\t__________________________")
    add_left_text(doc, "Project Coordinator/Co-coordinator\t\t\t\tHead of Department")

    doc.add_page_break()

    # --- Page 2 ---
    add_centered_text(doc, "Noida Institute of Engineering and Technology, Greater Noida\n(An Autonomous Institute)", size=14, bold=True)
    add_centered_text(doc, "School of Computer Science & Information Technology", size=12, bold=True)
    add_centered_text(doc, "\nPROJECT PROGRESS TRACKING RECORD\n", size=13, bold=True)

    # --- Table ---
    table = doc.add_table(rows=11, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Meeting\nNo.", "Meeting Date & Time", "Work Assigned by Supervisor", "Work Completed by Student", "Supervisor’s Remarks", "Supervisor’s Signature"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    meetings = [
        ("1", "Week 1", "Establish ML baseline using Random Forest on UCI dataset.", "Ingested data, scaled features, and built RF baseline model.", "Good baseline. Switch to XGBoost next week for better tabular performance.", ""),
        ("2", "Week 2", "Transition pipeline to XGBoost architecture.", "Integrated XGBoost. Compared MAE and RMSE metrics against Random Forest.", "Excellent transition. Focus entirely on Hyperparameter tuning via GridSearchCV next.", ""),
        ("3", "Week 3", "Perform Hyperparameter tuning to maximize accuracy.", "Implemented GridSearchCV (n_estimators, max_depth). Error reduced significantly.", "Great improvement. Freeze model weights. Begin Streamlit UI development.", ""),
        ("4", "", "", "", "", ""),
        ("5", "", "", "", "", ""),
        ("6", "", "", "", "", ""),
        ("7", "", "", "", "", ""),
        ("8", "", "", "", "", ""),
        ("9", "", "", "", "", ""),
        ("10", "", "", "", "", "")
    ]

    for row_idx, meeting in enumerate(meetings, start=1):
        for col_idx, text in enumerate(meeting):
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("\n")
    add_left_text(doc, "Remark by Project Coordinator/Co-coordinator: ____________________________________")

    # Save Document
    doc.save("Annexure3.docx")
    print(f"Generated successfully: Annexure3.docx")

if __name__ == "__main__":
    create_annexure3()
